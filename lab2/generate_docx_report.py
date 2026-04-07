from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
WEATHER_FILE = ROOT / "weather.pl"
PROJECT_FILE = ROOT / "project_car_search.pl"
OUTPUT_FILE = ROOT / "lab2_report_ready.docx"


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_title(doc: Document) -> None:
    p = doc.add_paragraph("ОТЧЕТ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("по лабораторной работе №2")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("Тема: «Введение в SWI-Prolog. Экспертная система \"Погода и рекомендации\"»")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Студент: Мухаметзянов Малик")
    doc.add_paragraph("Группа: ____________")
    doc.add_paragraph("Преподаватель: ____________")
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    heading.style.font.name = "Times New Roman"
    heading.style.font.size = Pt(14)


def add_text(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def main() -> None:
    weather_code = WEATHER_FILE.read_text(encoding="utf-8")
    project_code = PROJECT_FILE.read_text(encoding="utf-8")

    doc = Document()
    setup_document(doc)
    add_title(doc)

    add_heading(doc, "1. Цель работы")
    add_text(
        doc,
        "Изучить основы логического программирования на языке SWI-Prolog и реализовать экспертную систему на фактах и правилах.",
    )

    add_heading(doc, "2. Краткое описание программы")
    add_text(
        doc,
        "В файле weather.pl реализована экспертная система рекомендаций по погоде с учетом температуры, погодного состояния, ветра и активности.",
    )
    add_text(
        doc,
        "Также выполнен один проект на выбор из списка: №3 «Поиск нарушений по автомобилю» (файл project_car_search.pl).",
    )

    add_heading(doc, "3. Листинг кода с комментариями")
    add_heading(doc, "3.1. weather.pl", 2)
    add_code(doc, weather_code)
    add_heading(doc, "3.2. project_car_search.pl (проект №3)", 2)
    add_code(doc, project_code)

    add_heading(doc, "4. Результаты выполнения заданий")

    add_heading(doc, "4.1. Задание 1: добавление нового города", 2)
    add_text(doc, "Добавлен факт weather('Екатеринбург', -15, снег, 9).")
    add_code(doc, "?- show_all_advices.")
    add_code(
        doc,
        "Екатеринбург: Очень холодно, одевайтесь максимально тепло! Осторожно, снегопад, будьте аккуратны на дорогах. "
        "Ветрено, застегните куртку. Хорошая погода для катания на лыжах.",
    )

    add_heading(doc, "4.2. Задание 2: учет ветра", 2)
    add_code(doc, "?- advice_for('Санкт-Петербург', A).")
    add_code(
        doc,
        "A = 'Прохладно, возьмите ветровку или легкую куртку. Осторожно, снегопад, будьте аккуратны на дорогах. "
        "Сильный ветер, будьте осторожны! Можно просто погулять по городу.'",
    )

    add_heading(doc, "4.3. Задание 3: рекомендация активности", 2)
    add_code(
        doc,
        """?- activity_advice(27, солнечно, A1).
?- activity_advice(18, дождь, A2).
?- activity_advice(-5, снег, A3).
?- activity_advice(10, облачно, A4).""",
    )
    add_code(
        doc,
        """A1 = 'Идеально для пляжа!'
A2 = 'Можно посетить музей или кафе.'
A3 = 'Хорошая погода для катания на лыжах.'
A4 = 'Можно просто погулять по городу.'""",
    )

    add_heading(doc, "5. Выполнение 1 проекта на выбор")
    add_text(doc, "Выбран проект №3: «Поиск нарушений по автомобилю».")
    add_code(
        doc,
        """?- find_violations_by_car('А123ВС116', L).
L = [row('А123ВС116',speed_limit,202604011030),
     row('А123ВС116',red_light,202604031540)].

?- show_violations_by_car('А123ВС116').
Нарушения для А123ВС116:
- Тип: speed_limit, Время: 202604011030
- Тип: red_light, Время: 202604031540

?- show_violations_by_car('М001ТТ116').
Для автомобиля М001ТТ116 нарушений не найдено.""",
    )

    add_heading(doc, "6. Ответы на контрольные вопросы")
    add_text(doc, "1. Факт — безусловно истинное утверждение. Правило — условное утверждение вида Голова :- Тело.")
    add_text(doc, "2. Запрос с переменной: advice_for('Москва', A).")
    add_text(doc, "3. forall/2 нужен для выполнения действия для каждого решения условия.")
    add_text(doc, "4. Для объединения строк используется atomic_list_concat/3.")

    add_heading(doc, "7. Заключение")
    add_text(
        doc,
        "В работе реализованы все 3 обязательных задания и один проект на выбор (№3). "
        "Получена рабочая экспертная система и пример применения Prolog для задач поиска по базе фактов.",
    )

    add_heading(doc, "8. Скриншоты")
    add_text(
        doc,
        "Добавьте скриншоты выполнения запросов в SWISH: show_all_advices и запросы по project_car_search.pl.",
    )

    doc.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    main()
